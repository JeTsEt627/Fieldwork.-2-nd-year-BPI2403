"""Юнит-тесты парсинга документов (BE-04).

Используются минимальные PDF/DOCX, сгенерированные в памяти, чтобы не хранить
бинарные файлы в репозитории (кроме /tests/fixtures).
"""

import io
import os

import pytest
from docx import Document as DocxDocument

from app.core.exceptions import DocumentParsingError, EmptyDocumentError
from app.services.document_parser import ParsedPage, extract_text, normalize_text

# --------------------------------------------------------------------------- #
# Вспомогательные фабрики                                                     #
# --------------------------------------------------------------------------- #

def make_docx_bytes(paragraphs: list[str]) -> bytes:
    """Создать DOCX в памяти с заданными абзацами."""
    doc = DocxDocument()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_empty_docx_bytes() -> bytes:
    """Создать пустой DOCX без текста."""
    doc = DocxDocument()
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --------------------------------------------------------------------------- #
# normalize_text                                                               #
# --------------------------------------------------------------------------- #

class TestNormalizeText:
    def test_collapses_whitespace(self):
        assert normalize_text("hello   world") == "hello world"

    def test_strips_leading_trailing(self):
        assert normalize_text("  text  ") == "text"

    def test_collapses_newlines(self):
        assert normalize_text("line1\n\nline2") == "line1 line2"

    def test_empty_string_returns_empty(self):
        assert normalize_text("") == ""

    def test_whitespace_only_returns_empty(self):
        assert normalize_text("   \n\t  ") == ""

    def test_preserves_content(self):
        assert normalize_text("Машинное обучение") == "Машинное обучение"


# --------------------------------------------------------------------------- #
# DOCX-парсер                                                                 #
# --------------------------------------------------------------------------- #

class TestExtractDocx:
    def test_extracts_paragraphs(self):
        content = "Первый абзац. Второй абзац."
        file_bytes = make_docx_bytes([content])
        pages = extract_text(file_bytes, "docx")
        assert len(pages) == 1
        assert isinstance(pages[0], ParsedPage)
        assert "Первый абзац" in pages[0].text

    def test_multiple_paragraphs_merged(self):
        file_bytes = make_docx_bytes(["Абзац первый.", "Абзац второй."])
        pages = extract_text(file_bytes, "docx")
        assert len(pages) == 1
        assert "Абзац первый" in pages[0].text
        assert "Абзац второй" in pages[0].text

    def test_page_number_is_one(self):
        file_bytes = make_docx_bytes(["Текст документа."])
        pages = extract_text(file_bytes, "docx")
        assert pages[0].page_number == 1

    def test_empty_docx_raises_empty_document_error(self):
        file_bytes = make_empty_docx_bytes()
        with pytest.raises(EmptyDocumentError):
            extract_text(file_bytes, "docx")

    def test_invalid_bytes_raises_parsing_error(self):
        with pytest.raises(DocumentParsingError):
            extract_text(b"not a docx file at all", "docx")

    def test_case_insensitive_extension(self):
        file_bytes = make_docx_bytes(["Текст."])
        pages = extract_text(file_bytes, "DOCX")
        assert len(pages) == 1

    def test_extension_with_dot(self):
        file_bytes = make_docx_bytes(["Текст."])
        pages = extract_text(file_bytes, ".docx")
        assert len(pages) == 1


# --------------------------------------------------------------------------- #
# Неподдерживаемый формат                                                     #
# --------------------------------------------------------------------------- #

class TestExtractUnsupportedFormat:
    def test_raises_for_txt(self):
        with pytest.raises(DocumentParsingError):
            extract_text(b"plain text", "txt")

    def test_raises_for_empty_extension(self):
        with pytest.raises(DocumentParsingError):
            extract_text(b"data", "")


# --------------------------------------------------------------------------- #
# PDF-парсер (интеграционный, требует реального pdfplumber)                   #
# --------------------------------------------------------------------------- #

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")


@pytest.mark.skipif(
    not os.path.isdir(FIXTURES_DIR),
    reason="Директория tests/fixtures отсутствует",
)
class TestExtractPdfFromFixtures:
    def test_sample_pdf_parsed(self):
        # Берём только заведомо корректные PDF (valid_*.pdf); пустые и битые
        # файлы из набора QA-03 проверяются отдельными тестами.
        pdf_files = [
            f
            for f in os.listdir(FIXTURES_DIR)
            if f.lower().endswith(".pdf") and f.lower().startswith("valid")
        ]
        if not pdf_files:
            pytest.skip("Нет корректных PDF-файлов в tests/fixtures")
        with open(os.path.join(FIXTURES_DIR, pdf_files[0]), "rb") as fh:
            file_bytes = fh.read()
        pages = extract_text(file_bytes, "pdf")
        assert len(pages) > 0
        assert all(isinstance(p, ParsedPage) for p in pages)
        assert all(p.text for p in pages)


class TestExtractPdfInvalidBytes:
    def test_invalid_pdf_raises_parsing_error(self):
        with pytest.raises(DocumentParsingError):
            extract_text(b"definitely not a pdf", "pdf")
