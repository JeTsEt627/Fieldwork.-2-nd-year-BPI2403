"""Генератор набора тестовых документов (QA-03).

Создаёт в текущей директории (``backend/tests/fixtures``) полный набор файлов
для тестирования парсинга и валидации:

- корректные PDF и DOCX с текстом;
- пустые файлы (0 байт и валидный DOCX без текста);
- файлы с битым форматированием (мусорные байты под видом PDF/DOCX);
- PDF с нестандартным (встроенным) шрифтом.

DOCX-файлы создаются через ``python-docx`` (уже в зависимостях). PDF — через
``reportlab``; если он не установлен, генерация PDF пропускается с подсказкой.

Запуск:
    cd backend/tests/fixtures
    python generate_fixtures.py
"""

import os
import random

from docx import Document as DocxDocument

FIXTURES_DIR = os.path.dirname(os.path.abspath(__file__))

SAMPLE_TEXT = (
    "Машинное обучение — это раздел искусственного интеллекта, изучающий "
    "методы построения алгоритмов, способных обучаться на данных. "
    "Полнотекстовый поиск позволяет находить релевантные фрагменты документов "
    "по запросу пользователя с учётом морфологии русского языка."
)


def make_valid_docx() -> None:
    """Создать корректный DOCX с несколькими абзацами и таблицей."""
    doc = DocxDocument()
    doc.add_heading("Тестовый документ", level=1)
    doc.add_paragraph(SAMPLE_TEXT)
    doc.add_paragraph(
        "Второй абзац описывает индексацию документов в Elasticsearch и "
        "разбиение текста на чанки фиксированного размера."
    )
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Параметр"
    table.cell(0, 1).text = "Значение"
    table.cell(1, 0).text = "Размер чанка"
    table.cell(1, 1).text = "1000 символов"
    doc.save(os.path.join(FIXTURES_DIR, "valid_document.docx"))


def make_empty_docx() -> None:
    """Создать валидный DOCX без извлекаемого текста."""
    doc = DocxDocument()
    doc.save(os.path.join(FIXTURES_DIR, "empty_document.docx"))


def make_empty_pdf() -> None:
    """Создать пустой файл с расширением .pdf (0 байт)."""
    open(os.path.join(FIXTURES_DIR, "empty_file.pdf"), "wb").close()


def make_corrupted_files() -> None:
    """Создать файлы с битым форматированием (мусорные байты)."""
    random.seed(42)
    garbage = bytes(random.randint(0, 255) for _ in range(2048))

    # Поддельный PDF: корректная сигнатура, но повреждённое тело.
    with open(os.path.join(FIXTURES_DIR, "corrupted.pdf"), "wb") as fh:
        fh.write(b"%PDF-1.4\n" + garbage)

    # Поддельный DOCX: расширение docx, но содержимое не ZIP-архив.
    with open(os.path.join(FIXTURES_DIR, "corrupted.docx"), "wb") as fh:
        fh.write(garbage)


def make_pdf_files() -> bool:
    """Создать корректный PDF и PDF с нестандартным шрифтом.

    Returns:
        ``True``, если PDF-файлы созданы; ``False``, если ``reportlab`` не
        установлен.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
    except ImportError:
        return False

    # Корректный PDF со стандартным шрифтом.
    valid_path = os.path.join(FIXTURES_DIR, "valid_document.pdf")
    pdf = canvas.Canvas(valid_path, pagesize=A4)
    pdf.setFont("Helvetica", 12)
    text_obj = pdf.beginText(50, 800)
    for line in _wrap(SAMPLE_TEXT, 70):
        text_obj.textLine(line)
    pdf.drawText(text_obj)
    pdf.showPage()
    pdf.save()

    # PDF с нестандартным (моноширинным) встроенным шрифтом.
    font_path = os.path.join(FIXTURES_DIR, "unusual_font.pdf")
    pdf = canvas.Canvas(font_path, pagesize=A4)
    pdf.setFont("Courier-Oblique", 11)
    text_obj = pdf.beginText(50, 800)
    for line in _wrap("Nonstandard font sample. " + SAMPLE_TEXT, 70):
        text_obj.textLine(line)
    pdf.drawText(text_obj)
    pdf.showPage()
    pdf.save()
    return True


def _wrap(text: str, width: int) -> list[str]:
    """Разбить строку на строки заданной ширины (упрощённый перенос)."""
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        if len(current) + len(word) + 1 > width:
            lines.append(current)
            current = word
        else:
            current = f"{current} {word}".strip()
    if current:
        lines.append(current)
    return lines


def main() -> None:
    """Сгенерировать весь набор фикстур."""
    make_valid_docx()
    make_empty_docx()
    make_empty_pdf()
    make_corrupted_files()
    pdf_ok = make_pdf_files()

    print(f"Фикстуры созданы в: {FIXTURES_DIR}")
    print("  - valid_document.docx (корректный DOCX)")
    print("  - empty_document.docx (DOCX без текста)")
    print("  - empty_file.pdf (пустой файл)")
    print("  - corrupted.pdf, corrupted.docx (битое форматирование)")
    if pdf_ok:
        print("  - valid_document.pdf (корректный PDF)")
        print("  - unusual_font.pdf (нестандартный шрифт)")
    else:
        print("  ! PDF пропущены: установите reportlab "
              "(pip install reportlab) и запустите снова.")


if __name__ == "__main__":
    main()
