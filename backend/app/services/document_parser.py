"""Извлечение текста из загруженных документов (BE-04).

Поддерживаются форматы PDF (через ``pdfplumber``) и DOCX (через
``python-docx``). Текст возвращается постранично, что позволяет сохранять
номер страницы для каждого чанка при индексации (BE-07).
"""

import io
import re
from dataclasses import dataclass

import pdfplumber
from docx import Document as DocxDocument

from app.core.exceptions import DocumentParsingError, EmptyDocumentError
from app.core.logging_config import get_logger

logger = get_logger(__name__)

# Множество расширений, обрабатываемых парсером.
PDF_EXTENSION = "pdf"
DOCX_EXTENSION = "docx"


@dataclass(frozen=True)
class ParsedPage:
    """Текст одной страницы документа.

    Attributes:
        page_number: Номер страницы (начиная с 1).
        text: Извлечённый и нормализованный текст страницы.
    """

    page_number: int
    text: str


def normalize_text(text: str) -> str:
    """Нормализовать извлечённый текст.

    Убирает пустые строки, схлопывает последовательности пробельных символов
    в один пробел и удаляет ведущие/замыкающие пробелы. Это делает чанки
    компактнее и улучшает качество полнотекстового поиска.

    Args:
        text: Исходный текст.

    Returns:
        Нормализованный текст.
    """
    if not text:
        return ""
    # Заменяем любые последовательности пробельных символов одним пробелом.
    collapsed = re.sub(r"\s+", " ", text)
    return collapsed.strip()


def _extract_pdf(file_bytes: bytes) -> list[ParsedPage]:
    """Извлечь текст из PDF постранично с помощью pdfplumber.

    Args:
        file_bytes: Содержимое PDF-файла.

    Returns:
        Список непустых страниц с текстом.

    Raises:
        DocumentParsingError: Если файл не удаётся прочитать как PDF.
    """
    pages: list[ParsedPage] = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                raw_text = page.extract_text() or ""
                normalized = normalize_text(raw_text)
                if normalized:
                    pages.append(ParsedPage(page_number=index, text=normalized))
    except DocumentParsingError:
        raise
    except Exception as exc:  # pdfplumber/pdfminer бросают разнородные ошибки
        logger.exception("Не удалось разобрать PDF-файл")
        raise DocumentParsingError(f"Ошибка чтения PDF: {exc}") from exc
    return pages


def _extract_docx(file_bytes: bytes) -> list[ParsedPage]:
    """Извлечь текст из DOCX с помощью python-docx.

    Формат DOCX не хранит сведения о разбиении на страницы, поэтому весь
    документ возвращается как одна логическая «страница» (page_number=1).

    Args:
        file_bytes: Содержимое DOCX-файла.

    Returns:
        Список из одной страницы с текстом (или пустой список).

    Raises:
        DocumentParsingError: Если файл не удаётся прочитать как DOCX.
    """
    try:
        document = DocxDocument(io.BytesIO(file_bytes))
    except Exception as exc:
        logger.exception("Не удалось разобрать DOCX-файл")
        raise DocumentParsingError(f"Ошибка чтения DOCX: {exc}") from exc

    paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]

    # Дополнительно извлекаем текст из таблиц.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    paragraphs.append(cell.text)

    full_text = normalize_text("\n".join(paragraphs))
    if not full_text:
        return []
    return [ParsedPage(page_number=1, text=full_text)]


def extract_text(file_bytes: bytes, extension: str) -> list[ParsedPage]:
    """Извлечь текст из документа в зависимости от его формата (BE-04).

    Args:
        file_bytes: Содержимое файла.
        extension: Расширение файла без точки в нижнем регистре (``pdf``/``docx``).

    Returns:
        Список страниц с непустым текстом.

    Raises:
        DocumentParsingError: Формат не поддерживается или файл повреждён.
        EmptyDocumentError: В документе не найдено извлекаемого текста.
    """
    extension = extension.lower().lstrip(".")

    if extension == PDF_EXTENSION:
        pages = _extract_pdf(file_bytes)
    elif extension == DOCX_EXTENSION:
        pages = _extract_docx(file_bytes)
    else:
        raise DocumentParsingError(f"Неподдерживаемый формат: {extension}")

    if not pages:
        raise EmptyDocumentError("Документ не содержит извлекаемого текста")

    return pages
